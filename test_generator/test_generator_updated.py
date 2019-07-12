# This is the updated version that supports subscripts and superscripts. In order to use this file you must also use
# simple_xlsx.py which is uploaded to this repository. This addition was completed by Reddit user _DTR_.
# Thanks so much for the help on this!

import random
import platform
import re
import os
import sys
import tkinter as tk
import PySimpleGUI as sg
from docx import Document
from docx.shared import Inches

from simple_xlsx import load_workbook
from simple_xlsx import CellHelpers


def set_icon(app):
    if platform.system() == "Windows":
        app.iconbitmap(resource_path(os.path.join("Assets", "icon.ico")))
    elif platform.system() == "Linux":
        imgicon = tk.PhotoImage(file=resource_path(os.path.join("Assets", "icon.gif")))
        app.tk.call("wm", "iconphoto", app._w, imgicon)


def resource_path(relative_path):
    """ Get absolute path to resource, works for dev and for PyInstaller """
    try:
        # PyInstaller creates a temp folder and stores path in _MEIPASS
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)


def display_license():
    sg.PopupOK('License', 'The MIT License (MIT)\n'
                '=====================\n\n'
                'Copyright (c) 2019 Thomas Kellough\n\n'
                'Permission is hereby granted, free of charge, to any person obtaining a copy of\n'
                'this software and associated documentation files (the "Software"), to deal in\n'
                'the Software without restriction, including without limitation the rights to\n'
                'use, copy, modify, merge, publish, distribute, sublicense, and/or sell copies\n'
                'of the Software, and to permit persons to whom the Software is furnished to do\n'
                'so, subject to the following conditions:\n\n'

                'The above copyright notice and this permission notice shall be included in all '
                'copies or substantial portions of the Software.\n\n'

                'THE SOFTWARE IS PROVIDED "AS IS", WITHOUT WARRANTY OF ANY KIND, EXPRESS OR\n'
                'IMPLIED, INCLUDING BUT NOT LIMITED TO THE WARRANTIES OF MERCHANTABILITY,\n'
                'FITNESS FOR A PARTICULAR PURPOSE AND NONINFRINGEMENT. IN NO EVENT SHALL THE\n'
                'AUTHORS OR COPYRIGHT HOLDERS BE LIABLE FOR ANY CLAIM, DAMAGES OR OTHER\n'
                'LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR OTHERWISE, ARISING FROM,\n'
                'OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE OR OTHER DEALINGS IN THE\n'
                'SOFTWARE.')


path = os.path.join(os.path.abspath('test_bank.xlsx'))
image_path = os.path.join(os.path.abspath('Images'))
image_regex = re.compile(r'\(.*\)')

# Used to write the answer key
# The index is found from the correct answer dictionary and matched to the correct answer choice
index_choice_match = {0: "A",
                      1: "B",
                      2: "C",
                      3: "D",
                      4: "E"}


def main():
    class MyTestGenerator:
        wb_obj = load_workbook(path)
        sheet_obj = wb_obj.active
        max_col = sheet_obj.dim['col_last']
        max_row = sheet_obj.dim['rw_last']

        def __init__(self):
            self.unique_unit_list = []
            self.questions_and_answers = {}  # Includes all questions and ALL answer choices for each test
            self.current_unit_list = []
            self.previous_units = []
            self.not_shuffled_answer_choices = []
            self.question_and_correct_answer = {}
            self.correct_answer_key = []

        def create_random_test(self, current_unit, current_unit_num_questions, previous_unit_num_questions, num_of_answer_choices):
            # Create unique value lists based off user selection (current unit & number of questions)
            # We only have one column, grab the first and only one
            for cell in self.sheet_obj.get_range('C2:C' + str(self.max_row)):
                unique_id = int(self.sheet_obj.cell(cell.rw, col=1).value.plain_text()) # Units are numbers. Strip formatting and convert to int
                if cell.value.plain_text() != current_unit:
                    self.previous_units.append(unique_id)
                else:
                    self.current_unit_list.append(unique_id)

            current_unit_unique_id = random.sample(self.current_unit_list, int(current_unit_num_questions))
            previous_unit_unique_id = random.sample(self.previous_units, int(previous_unit_num_questions))
            all_questions_unique_id = previous_unit_unique_id + current_unit_unique_id

            # Create dictionary of questions and answer choices from both unique lists
            for cell in self.sheet_obj.get_range('A2:A' + str(self.max_row)):
                if int(self.sheet_obj.cell(cell.rw, col=1).value.plain_text()) not in all_questions_unique_id:
                    continue

                question = self.sheet_obj.cell(cell.rw, col=4).value
                for occurrence in [' I.', ' II.', ' III.', ' IV.']:
                    question.replace(occurrence, '\n' + occurrence)

                answer_choice_range = CellHelpers.build_range(cell.rw, cell.rw, 5, 5 + num_of_answer_choices - 1)
                answer_choices = [answer.value for answer in self.sheet_obj.get_range(answer_choice_range)]

                self.questions_and_answers.update({question: answer_choices})
                self.question_and_correct_answer.update({question: answer_choices[0]})

        def write_test(self, filename, num_of_answer_keys, num_of_answer_choices):
            for answer_key in num_of_answer_keys:
                filename = f'{filename} - {answer_key}'
                doc = Document()
                doc.add_paragraph(f'Name:\nDate:')
                doc.add_paragraph(f'{filename}').style = 'Title'

                # Randomly write both test questions and answers from dictionary
                for count, (q, answer_choices) in enumerate(sorted(self.questions_and_answers.items(), key=lambda x: random.random())):
                    question = q
                    mo = image_regex.search(question.plain_text())
                    shuffled = answer_choices
                    self.not_shuffled_answer_choices.append(answer_choices[0])
                    random.shuffle(answer_choices)

                    # Create answer key for each test
                    correct_answer_index = answer_choices.index(self.question_and_correct_answer[question])
                    self.correct_answer_key.append(index_choice_match[correct_answer_index])

                    # If question needs image, add from image directory
                    if mo:
                        string_regex = mo.group().replace('(', '')
                        string_regex = string_regex.replace(')', '')
                        string_regex = f'{string_regex}.png'
                        doc.add_picture(f'{image_path}\\{string_regex}', width=Inches(4.0))
                        question.replace(mo.group(), "(Use the above figure to help with this question)")

                    paragraph = doc.add_paragraph()
                    question.add_to_paragraph(paragraph)
                    paragraph.add_run('\n')
                    paragraph.style = 'List Number'
                    if mo:
                        question.replace("(Use the above figure to help with this question)", mo.group())

                    # Randomly write answer choices for each question & create question/answer dictionary
                    for choice, answer in enumerate(num_of_answer_choices):
                        paragraph.add_run(f'\t{answer}) ')
                        shuffled[choice].add_to_paragraph(paragraph)
                        paragraph.add_run('\n')

                doc.add_paragraph('\nAnswer Key\n\n')
                for i in self.correct_answer_key:
                    paragraph_answer_key = doc.add_paragraph(i)
                    paragraph_answer_key.style = "List Number"
                self.correct_answer_key.clear()
                doc.save(f'{filename}.docx')
                filename = filename.replace(f' - {answer_key}', '').strip()

        # Creates list for user to select the current unit
        def create_unique_unit_list(self):
            for row in range(2, self.max_row):
                unit = self.sheet_obj.cell(rw=row, col=3).value.plain_text()
                self.unique_unit_list.append(unit)
            self.unique_unit_list = list(dict.fromkeys(self.unique_unit_list))
            return self.unique_unit_list

    p = MyTestGenerator()
    sg.ChangeLookAndFeel('Reddit')
    menu_def = [['File', ['Exit']],
                ['Help', 'License']]

    layout = [
        [sg.Menu(menu_def)],
        [sg.Text('Random Test Generator', font=('Helvatica', 20), text_color='MidnightBlue')],
        [sg.Text('Select Current Unit'), sg.Listbox(values=p.create_unique_unit_list(), size=(30, 5))],
        [sg.Text('How many questions from the current unit?'),
         sg.Slider(range=(0, 100), orientation='h', size=(34, 20), default_value=15)],
        [sg.Text('How many questions from the previous unit?'),
         sg.Slider(range=(0, 100), orientation='h', size=(34, 20), default_value=35)],
        [sg.Text('Name of your test: '), sg.InputText()],
        [sg.Text('Number of keys: '), sg.Spin(values=(1, 2, 3, 4), initial_value=1), sg.Text('Number of answer options'),
         sg.Spin(values=(3, 4, 5), initial_value=4)],
        [sg.Button('Create Test'), sg.Quit()]
       ]

    window = sg.Window('Random Test Generator').Layout(layout)
    window.SetIcon(os.path.join(os.path.abspath('assets\\icon.ico')))

    while True:
        event, values = window.Read()

        # User inputs from GUI
        current_unit = values[1][0]
        current_unit_num_questions = values[2]
        previous_unit_num_questions = values[3]
        filename = values[4]
        num_of_answer_keys = values[5]
        num_of_answer_choices = values[6]

        key_list = {1: ['A'], 2: ['A', 'B'], 3: ['A', 'B', 'C'], 4: ['A', 'B', 'C', 'D']}
        answer_choice_options = {3: ['A', 'B', 'C'], 4: ['A', 'B', 'C', 'D'], 5: ['A', 'B', 'C', 'D', 'E']}

        if event == 'Create Test':
            try:
                if filename == '':
                    sg.PopupOK('Please enter a name for your file')
                    continue
                else:
                    p.create_random_test(current_unit=current_unit, current_unit_num_questions=current_unit_num_questions,
                                         previous_unit_num_questions=previous_unit_num_questions,
                                         num_of_answer_choices=int(num_of_answer_choices))
                    p.write_test(filename=filename, num_of_answer_keys=key_list[int(num_of_answer_keys)],
                                 num_of_answer_choices=answer_choice_options[int(num_of_answer_choices)])
                    sg.PopupOK('Test created!')
                    break
            except IndexError:
                sg.PopupOK('Error!', 'Please select your current unit')
            except ValueError:
                sg.PopupOK('Error!', f'Question numbers out of range.')
        elif event == 'License':
            display_license()
        elif event == 'Quit' or event == 'Exit' or event is None:
            window.Close()
            break


if __name__ == '__main__':
    main()
