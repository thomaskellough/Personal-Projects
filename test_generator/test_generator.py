import openpyxl
import random
import docx
import re
import os
import PySimpleGUI as sg
from docx.shared import Inches

path = os.path.join(os.path.abspath('test_bank.xlsx'))
image_path = os.path.join(os.path.abspath('Images'))
image_regex = re.compile(r'\(Ch_\d\d_\d\d\)')


def display_license():
    sg.PopupOK('License', 'The MIT License (MIT)\n'
                '=====================\n\n'
                'Copyright (c) 2019 Thomas Kellough\n\n'
                'Permission is hereby granted, free of charge, to any person obtaining a copy of\n'
                'this software and associated documentation files (the "Software"), to deal in\n'
                'the Software without restriction, including without limitation the rights to\n'
                'use, copy, modify, merge, publish, distribute, sublicense, and/or sell copies\n'
                'of the Software, and to permit persons to whom the Software is furnished to do\n'
                'so, subject to the following conditions:\n\n'

                'The above copyright notice and this permission notice shall be included in all '
                'copies or substantial portions of the Software.\n\n'

                'THE SOFTWARE IS PROVIDED "AS IS", WITHOUT WARRANTY OF ANY KIND, EXPRESS OR\n'
                'IMPLIED, INCLUDING BUT NOT LIMITED TO THE WARRANTIES OF MERCHANTABILITY,\n'
                'FITNESS FOR A PARTICULAR PURPOSE AND NONINFRINGEMENT. IN NO EVENT SHALL THE\n'
                'AUTHORS OR COPYRIGHT HOLDERS BE LIABLE FOR ANY CLAIM, DAMAGES OR OTHER\n'
                'LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR OTHERWISE, ARISING FROM,\n'
                'OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE OR OTHER DEALINGS IN THE\n'
                'SOFTWARE.')


class MyTestGenerator:
    wb_obj = openpyxl.load_workbook(path)
    sheet_obj = wb_obj.active
    max_col = sheet_obj.max_column
    max_row = sheet_obj.max_row

    def __init__(self):
        self.unique_unit_list = []
        self.questions_and_answers = {}
        self.chapter_distribution = {}
        self.test_key = {}
        self.current_unit_list = []
        self.previous_unit = []
        self.not_shuffled_answer_choices = []

    def create_random_test(self, current_unit, current_unit_num_questions, previous_unit_num_questions):
        current_unit_num_questions = int(current_unit_num_questions)
        previous_unit_num_questions = int(previous_unit_num_questions)
        current_unit = current_unit
        for row in range(2, self.max_row + 1):
            if self.sheet_obj.cell(row=row, column=3).value != current_unit:
                self.previous_unit.append(self.sheet_obj.cell(row=row, column=1).value)
            else:
                self.current_unit_list.append(self.sheet_obj.cell(row=row, column=1).value)

        current_unit_unique_id = random.sample(self.current_unit_list, current_unit_num_questions)
        previous_unit_unique_id = random.sample(self.previous_unit, previous_unit_num_questions)
        all_questions_unique_id = previous_unit_unique_id + current_unit_unique_id
        for row in all_questions_unique_id:
            unit = self.sheet_obj.cell(row=row, column=2).value
            if unit not in self.chapter_distribution.keys():
                self.chapter_distribution[unit] = 1
            else:
                self.chapter_distribution[unit] += 1
        for row in range(2, self.max_row + 1):
            answer_choices = []
            if self.sheet_obj.cell(row=row, column=1).value in all_questions_unique_id:
                question = self.sheet_obj.cell(row=row, column=4).value

                for occurrence in [' I.', ' II.', ' III.', ' IV.']:
                    question = question.replace(occurrence, '\n' + occurrence)

                for column in range(5):
                    answer = self.sheet_obj.cell(row=row, column=column + 5).value
                    answer_choices.append(answer)
                self.questions_and_answers.update({question: answer_choices})

    def write_test(self, filename):
        if filename == '':
            filename = 'New_Test'
        doc = docx.Document()
        doc.add_paragraph(f'Name:\nDate:')
        doc.add_paragraph(f'{filename}').style = 'Title'

        for count, (key, value) in enumerate(self.questions_and_answers.items()):
            question = f'{key}'
            mo = image_regex.search(question)
            shuffled = value
            self.not_shuffled_answer_choices.append(value[0])
            random.shuffle(value)

            if mo:
                string_regex = mo.group().replace('(', '')
                string_regex = string_regex.replace(')', '')
                string_regex = f'{string_regex}.png'
                doc.add_picture(f'{image_path}\\{string_regex}', width=Inches(4.0))
                paragraph = doc.add_paragraph(f'{question.replace(mo.group(), "(Use the above figure to help with this question)")}\n')
            else:
                paragraph = doc.add_paragraph(f'{question}\n')
            paragraph.style = 'List Number'

            for choice, answer in enumerate(['A', 'B', 'C', 'D', 'E']):
                if shuffled[choice] == self.not_shuffled_answer_choices[count]:
                    self.test_key.update({count + 1: answer})
                paragraph_answers = f'\t{answer}) {shuffled[choice]}\n'
                paragraph.add_run(paragraph_answers)

        doc.add_paragraph('\nAnswer Key\n\n')

        for key, value in self.test_key.items():
            paragraph_answer_key = doc.add_paragraph(f'{value}')
            paragraph_answer_key.style = 'List Number'
        doc.save(f'{filename}.docx')

    def create_unique_unit_list(self):
        for row in range(2, self.max_row):
            unit = self.sheet_obj.cell(row=row, column=3).value
            self.unique_unit_list.append(unit)
        self.unique_unit_list = list(dict.fromkeys(self.unique_unit_list))
        return self.unique_unit_list


p = MyTestGenerator()
sg.ChangeLookAndFeel('Reddit')
menu_def = [['File', ['Exit']],
            ['Help', 'License'], ]

layout = [
    [sg.Menu(menu_def)],
    [sg.Text('Random Test Generator')],
    [sg.Text('Select Current Unit'), sg.Listbox(values=p.create_unique_unit_list(), size=(30, 5))],
    [sg.Text('How many questions from the current unit?'),
     sg.Slider(range=(1, 100), orientation='h', size=(34, 20), default_value=10)],
    [sg.Text('How many questions from the previous unit?'),
     sg.Slider(range=(1, 100), orientation='h', size=(34, 20), default_value=15)],
    [sg.Text('Name of your test: '), sg.InputText()],
    [sg.Button('Create Test'), sg.Quit()]
   ]

window = sg.Window('Random Test Generator').Layout(layout)
window.SetIcon(os.path.join(os.path.abspath('assets\\icon.ico')))

while True:
    event, values = window.Read()
    if event == 'Create Test':
        try:
            p.create_random_test(current_unit=values[1][0], current_unit_num_questions=values[2],
                                 previous_unit_num_questions=values[3])
            p.write_test(filename=values[4])
            sg.PopupOK('Test created!')
        except IndexError:
            sg.PopupOK('Error!', 'Please select your current unit')
        except ValueError:
            sg.PopupOK('Error!', f'Question numbers out of range.')
    elif event == 'License':
        display_license()
    elif event == 'Quit' or event == 'Exit' or event is None:
        window.Close()
        break
